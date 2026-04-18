from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from src.utils.logger import get_logger
from src import config

logger = get_logger(__name__)


@dataclass
class DocumentResult:

    text: str
    source: str
    num_pages: Optional[int] = None
    metadata: dict = field(default = dict)


def parse_txt(file_path: Path) -> DocumentResult:

    logger.info(f"Parsing TXT: {file_path.name}")

    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        logger.warning(f"UTF-8 failed for {file_path.name}, trying latin-1")
        text = file_path.read_text(encoding="latin-1")

    return DocumentResult(
        text=text.strip(),
        source=file_path.name,
        metadata={"type": "txt"},
    )


def parse_pdf(file_path: Path) -> DocumentResult:

    from PyPDF2 import PdfReader

    logger.info(f"Parsing PDF: {file_path.name}")

    reader = PdfReader(str(file_path))
    pages_text = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            pages_text.append(page_text)
        else:
            logger.warning(f"Page {i + 1} of {file_path.name}: no text extracted (possibly scanned)")

    full_text = "\n\n".join(pages_text)

    return DocumentResult(
        text=full_text.strip(),
        source=file_path.name,
        num_pages=len(reader.pages),
        metadata={
            "type": "pdf",
            "pages_with_text": len(pages_text),
            "total_pages": len(reader.pages),
        },
    )


def parse_docx(file_path: Path) -> DocumentResult:

    from docx import Document

    logger.info(f"Parsing DOCX: {file_path.name}")

    doc = Document(str(file_path))

    # Extract paragraph text
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Extract table text (tables often contain important structured data)
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)

    # Combine: paragraphs first, then tables
    all_text = "\n\n".join(paragraphs)
    if table_texts:
        all_text += "\n\n[TABLE DATA]\n" + "\n".join(table_texts)

    return DocumentResult(
        text=all_text.strip(),
        source=file_path.name,
        metadata={
            "type": "docx",
            "num_paragraphs": len(paragraphs),
            "num_tables": len(doc.tables),
        },
    )


def parse_document(file_path: Path) -> DocumentResult:

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # Check file size
    size_mb = file_path.stat().st_size / (1024 * 1024)
    if size_mb > config.MAX_FILE_SIZE_MB:
        raise ValueError(
            f"File {file_path.name} is {size_mb:.1f}MB, "
            f"exceeds limit of {config.MAX_FILE_SIZE_MB}MB"
        )

    ext = file_path.suffix.lower()

    if ext not in config.SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {config.SUPPORTED_EXTENSIONS}"
        )

    # Route to the right parser
    parsers = {
        ".txt": parse_txt,
        ".md": parse_txt,  # Markdown is just text for our purposes
        ".pdf": parse_pdf,
        ".docx": parse_docx,
    }

    try:
        result = parsers[ext](file_path)
    except Exception as e:
        raise RuntimeError(f"Failed to parse {file_path.name}: {e}") from e

    if not result.text:
        logger.warning(f"No text extracted from {file_path.name}")

    logger.info(
        f"Parsed {file_path.name}: {len(result.text)} chars extracted"
    )

    return result